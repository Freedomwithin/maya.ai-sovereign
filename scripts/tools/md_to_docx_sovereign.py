import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, "r") as f:
        lines = f.readlines()

    doc = Document()
    
    # --- Style Settings ---
    # (Note: Standard python-docx doesn't easily support dark mode, 
    # but we can use Sovereign-style headers)
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:], 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], 1)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], 2)
        elif line.startswith("- "):
            doc.add_paragraph(line, style='List Bullet')
        elif line.startswith("*"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("*", ""))
            run.italic = True
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Success! Document forged at: {docx_path}")

if __name__ == "__main__":
    md_file = "memories/updates/2026-04-02_Overnight_Sovereign_Synthesis.md"
    docx_file = "memories/updates/2026-04-02_Overnight_Sovereign_Synthesis.docx"
    convert_md_to_docx(md_file, docx_file)
